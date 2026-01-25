#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
第3回 関西みらいベンチャーアワード『みらいWay』 エントリーシート作成スクリプト
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_entry_sheet():
    """エントリーシートのWordファイルを作成"""
    
    # ドキュメントを作成
    doc = Document()
    
    # タイトル
    title = doc.add_heading('第3回 関西みらいベンチャーアワード『みらいWay』 エントリーシート', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本情報
    doc.add_heading('1. 会社名（個人の場合はお名前）', level=2)
    doc.add_paragraph('CarnivOS')
    
    doc.add_heading('2. 会社名（フリガナ）', level=2)
    doc.add_paragraph('カーニボス')
    
    doc.add_heading('3. 設立年月', level=2)
    doc.add_paragraph('2025年12月18日')
    
    doc.add_heading('4. 本社所在地', level=2)
    doc.add_paragraph('〒659-0063')
    doc.add_paragraph('兵庫県芦屋市宮川町7-17')
    
    doc.add_heading('5. 電話番号', level=2)
    doc.add_paragraph('09021946761')
    
    doc.add_heading('6. メールアドレス', level=2)
    doc.add_paragraph('sibuketu@gmail.com')
    
    doc.add_heading('7. 代表者名', level=2)
    doc.add_paragraph('須佐見隼人')
    
    doc.add_heading('8. 代表者（フリガナ）', level=2)
    doc.add_paragraph('スサミハヤト')
    
    doc.add_heading('9. 代表者生年月日', level=2)
    doc.add_paragraph('2006年1月1日')
    
    doc.add_heading('10. 法人・個人', level=2)
    doc.add_paragraph('個人')
    
    doc.add_heading('11. 法人設立予定年月', level=2)
    doc.add_paragraph('2026年12月1日')
    
    # 事業概要
    doc.add_heading('12. 事業の概要・PRポイント（500文字以内）', level=2)
    pr_text = """CarnivOSは、世界初のカーニボアダイエット特化型栄養管理アプリです。
9つの世界に0個の独自機能（ButcherSelect、AIチャット、Recovery Protocol等）を搭載し、
完全なブルーオーシャン市場を開拓します。
月額$9.99（約¥1,500/月）、年額$69.99（約¥10,500/年）のサブスクリプションモデルで、
急成長するヘルスケアアプリ市場（2024年の25億米ドルから2033年には133億9,000万米ドルに急成長、CAGR 20.5%）で成功を目指します。
プロトタイプは既に完成しており、リリース準備段階です。"""
    doc.add_paragraph(pr_text)
    
    doc.add_heading('13. 事業の目的・ビジョン（500文字以内）', level=2)
    vision_text = """世界一のCarnivoreアプリを目指し、カーニボアコミュニティの健康管理を支援します。
カーニボアダイエット実践者が直面している4つの主要な課題
（適切な栄養管理ツールの不在、専門的なアドバイスの不足、違反時のリカバリー方法の不明、個人最適化の不足）
を包括的に解決し、カーニボアコミュニティの健康管理を支援します。"""
    doc.add_paragraph(vision_text)
    
    doc.add_heading('14. 事業の新規性・革新性（500文字以内）', level=2)
    innovation_text = """CarnivOSは、9つの世界に0個の独自機能を搭載しています。
ButcherSelect（動物・部位選択UI）、AIチャット（専門家レベル）、
Recovery Protocol（自動リカバリープロトコル生成）、動的RDA（栄養素目標値の動的調整）、
100項目以上のプロファイル設定など、既存の健康アプリやカーニボア専門アプリにはない機能を提供します。
これにより、完全なブルーオーシャン市場を開拓し、競合優位性を確立します。"""
    doc.add_paragraph(innovation_text)
    
    doc.add_heading('15. 起業又は事業立ち上げに至る背景・経緯（500文字以内）', level=2)
    background_text = """2025年12月18日にアプリのアイデアを思いつきました。
カーニボアダイエット実践者として、適切な栄養管理ツールの不在を実感し、
世界初のカーニボアダイエット特化型栄養管理アプリの開発を決意しました。
個人開発（Cursor + AI活用）により、開発コストを最小化しながら高度な機能を実装し、
ブートストラップ（初期資金0円）からスタートしました。
プロトタイプは既に完成しており、リリース準備段階です。"""
    doc.add_paragraph(background_text)
    
    doc.add_heading('16. 事業の現況と今後の見込み（500文字以内）', level=2)
    future_text = """現在、プロトタイプは完成しており、リリース準備段階です。
短期目標（0-6ヶ月）は、リリースと初期ユーザー100人獲得です。
中期目標（6-12ヶ月）は、月間アクティブユーザー500人とiOSアプリリリースです。
長期目標（12ヶ月以降）は、月間アクティブユーザー1,000人以上と国際展開です。
SNS中心のマーケティング戦略により、カーニボアコミュニティとの交流、ビジュアルコンテンツ、ショート動画を展開します。"""
    doc.add_paragraph(future_text)
    
    doc.add_heading('17. お取引金融機関', level=2)
    doc.add_paragraph('その他（未定）')
    
    doc.add_heading('18. お取引支店', level=2)
    doc.add_paragraph('未定')
    
    doc.add_heading('19. 本アワードを知ったきっかけ', level=2)
    doc.add_paragraph('当社ホームページ')
    
    # ファイルを保存
    # ワークスペースの絶対パスを使用
    workspace_path = r'C:\Users\susam\Downloads\新しいフォルダー\docs\primal-logic-app\primal-logic-web'
    output_path = os.path.join(workspace_path, 'docs', 'ENTRY_SHEET_MIRAIWAY.docx')
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    
    print(f'エントリーシートを作成しました: {output_path}')
    return output_path

if __name__ == '__main__':
    try:
        # スクリプトのパスを取得（__file__が使えない場合の対策）
        try:
            script_path = os.path.abspath(__file__)
        except NameError:
            # __file__が使えない場合は、カレントディレクトリから推測
            script_path = os.path.abspath('scripts/create_entry_sheet.py')
        
        # スクリプトのディレクトリに移動
        script_dir = os.path.dirname(script_path)
        if os.path.exists(script_dir):
            os.chdir(script_dir)
        
        output_path = create_entry_sheet()
        print(f'\n✅ エントリーシートの作成が完了しました')
        print(f'📄 ファイルパス: {output_path}')
    except ImportError:
        print('❌ python-docxライブラリがインストールされていません')
        print('以下のコマンドでインストールしてください:')
        print('  pip install python-docx')
    except Exception as e:
        print(f'❌ エラーが発生しました: {e}')
        import traceback
        traceback.print_exc()
